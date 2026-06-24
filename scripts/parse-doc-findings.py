#!/usr/bin/env python3
"""
Parse highlighted DOCX transcript to extract human-coded findings.

Extracts:
  - Highlighted text passages (the transcript excerpt)
  - Highlight color → mapped to objective category
  - Associated comments (the human insight/theme)
  - +/− markers → facilitator or barrier
  - Interview ID from headings/structure

Output: JSON file with all parsed findings.

Usage:
  pip install python-docx lxml
  python scripts/parse-doc-findings.py
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree
except ImportError:
    print("Install dependencies: pip install python-docx lxml")
    sys.exit(1)

# ── Config ──────────────────────────────────────────────────────────────

DOC_PATH = Path(__file__).parent.parent / "INTERVIEW TRANSCRIPTS.docx"
OUTPUT_PATH = Path(__file__).parent.parent / "parsed-human-findings.json"

# Word highlight color indices → names
# See: https://docs.microsoft.com/en-us/openspecs/office_standards/ms-oi29500/
HIGHLIGHT_COLORS = {
    "yellow": "yellow",
    "green": "green",
    "cyan": "cyan",
    "magenta": "magenta",
    "blue": "blue",
    "red": "red",
    "darkBlue": "darkBlue",
    "darkCyan": "darkCyan",
    "darkGreen": "darkGreen",
    "darkMagenta": "darkMagenta",
    "darkRed": "darkRed",
    "darkYellow": "darkYellow",
    "darkGray": "darkGray",
    "lightGray": "lightGray",
    "black": "black",
}

# Namespaces for XML parsing
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


def get_highlight_color(run_element):
    """Extract highlight color from a run's XML element."""
    rPr = run_element.find('.//w:rPr', NSMAP)
    if rPr is None:
        return None
    highlight = rPr.find('w:highlight', NSMAP)
    if highlight is not None:
        return highlight.get(f'{{{NSMAP["w"]}}}val')
    # Also check shading (some docs use shading instead of highlight)
    shd = rPr.find('w:shd', NSMAP)
    if shd is not None:
        fill = shd.get(f'{{{NSMAP["w"]}}}fill')
        if fill and fill.lower() not in ('auto', 'ffffff', 'none'):
            return f"shading:{fill}"
    return None


# Common body/UI text colors — NOT highlight colors
BODY_TEXT_COLORS = {
    '000000', '1a1f2c', '4a5263', '8a929c', '333333', '666666',
    '555555', '444444', '222222', '111111', 'auto',
    '0e5c5c',  # app's teal accent — not a doc highlight
    'b5bbc4', 'b8456d', 'b8842a',  # app UI colors
}


def get_run_color(run_element):
    """Extract font color from a run's XML element (ignoring body text colors)."""
    rPr = run_element.find('.//w:rPr', NSMAP)
    if rPr is None:
        return None
    color = rPr.find('w:color', NSMAP)
    if color is not None:
        val = color.get(f'{{{NSMAP["w"]}}}val')
        if val and val.lower() not in BODY_TEXT_COLORS:
            return val
    return None


def extract_comments(doc_path):
    """Extract all comments from the DOCX file."""
    from zipfile import ZipFile

    comments = {}
    with ZipFile(doc_path, 'r') as zf:
        # Check if comments.xml exists
        if 'word/comments.xml' not in zf.namelist():
            print("  ⚠ No comments.xml found in document")
            return comments

        comments_xml = zf.read('word/comments.xml')
        tree = etree.fromstring(comments_xml)

        for comment in tree.findall('.//w:comment', NSMAP):
            comment_id = comment.get(f'{{{NSMAP["w"]}}}id')
            author = comment.get(f'{{{NSMAP["w"]}}}author', '')
            date = comment.get(f'{{{NSMAP["w"]}}}date', '')

            # Get all text within the comment
            texts = []
            for t in comment.iter(f'{{{NSMAP["w"]}}}t'):
                if t.text:
                    texts.append(t.text)
            text = ' '.join(texts).strip()

            if comment_id:
                comments[comment_id] = {
                    'id': comment_id,
                    'author': author,
                    'date': date,
                    'text': text,
                }

    print(f"  Found {len(comments)} comments in document")
    return comments


def build_interview_boundaries(doc):
    """
    Scan the document for Heading 1 paragraphs (D-001, P-001, S-001 etc.)
    and section headers (Doctors:, Patients:, Survivors:) to build an ordered
    list of (paragraph_index, interview_id, participant_type).
    """
    boundaries = []
    current_section_type = None

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ''

        # Section headers — underlined, e.g. "Doctors:", "Patients:", "Survivors:"
        text_lower = text.lower().rstrip(':').strip()
        if text_lower in ('doctors', 'patients', 'survivors'):
            current_section_type = {
                'doctors': 'doctor',
                'patients': 'patient',
                'survivors': 'survivor',
            }[text_lower]
            continue

        # Heading 1 entries: D-001, P-001, S-001 etc.
        if 'Heading' in style_name:
            iid = text.strip()
            # Infer type from prefix if section wasn't set
            ptype = current_section_type
            if iid.startswith('D-'):
                ptype = 'doctor'
            elif iid.startswith('P-'):
                ptype = 'patient'
            elif iid.startswith('S-'):
                ptype = 'survivor'
            boundaries.append((idx, iid, ptype or 'other'))

    return boundaries


def get_interview_for_paragraph(para_idx, boundaries):
    """Given a paragraph index, find the interview it belongs to."""
    interview_id = None
    participant_type = None
    for bnd_idx, iid, ptype in boundaries:
        if bnd_idx <= para_idx:
            interview_id = iid
            participant_type = ptype
        else:
            break
    return interview_id, participant_type


def detect_facilitator_barrier(text):
    """Detect +/− markers in text to classify as facilitator or barrier."""
    text_stripped = text.strip()
    # Check end of text
    if text_stripped.endswith('(+)') or text_stripped.endswith('(+) '):
        return 'facilitator'
    if text_stripped.endswith('(-)') or text_stripped.endswith('(-) '):
        return 'barrier'
    # Check for + or - at the very end
    if text_stripped.endswith('+'):
        return 'facilitator'
    if text_stripped.endswith('-') and not text_stripped.endswith('--'):
        return 'barrier'
    # Check anywhere in last 10 chars
    tail = text_stripped[-10:]
    if '(+)' in tail:
        return 'facilitator'
    if '(-)' in tail:
        return 'barrier'
    return None


def map_color_to_objective(color):
    """Map highlight/font color to an objective category."""
    if not color:
        return None

    color_lower = color.lower()

    # Color mapping:
    #   Yellow / light yellow → Objective 1 (Early Detection)
    #   Green / light green  → Objective 2 (Diagnosis & Treatment)
    #   Blue  / light blue   → Objective 3 (Continuity & Follow-Up)

    # Shading colors (hex)
    if color_lower.startswith('shading:'):
        hex_color = color_lower.replace('shading:', '')
        # Yellow / peach family → Objective 1 (Early Detection)
        if hex_color in ('ffff00', 'ffff99', 'ffc000', 'fff2cc', 'ffe599', 'ffd966',
                         'f9cb9c', 'fce5cd', 'fff9e6', 'ffffcc', 'ffff66',
                         'f6b26b', 'e69138', 'b45f06'):
            return 'objective_1'
        # Green family → Objective 2 (Diagnosis & Treatment)
        if hex_color in ('00ff00', '00cc00', '008000', '70ad47', '00b050', '93c47d',
                         'b6d7a8', 'd9ead3', 'ccffcc', '99ff99',
                         '6aa84f', '38761d', '274e13',
                         'd0e0e3'):
            return 'objective_2'
        # Blue family → Objective 3 (Continuity & Follow-Up)
        if hex_color in ('0000ff', '0000cc', '0033ff', '3333ff', '4472c4', '6fa8dc',
                         '9fc5e8', 'cfe2f3', '3d85c6', '1155cc', '1c4587',
                         'ccccff', '9999ff', '6666ff', 'a4c2f4', '3c78d8',
                         'b4c6e7', 'd6e4f0', 'dbe5f1',
                         'c9daf8'):
            return 'objective_3'
        return f'unknown_color:{hex_color}'

    # Named highlight colors (Word built-in)
    COLOR_MAP = {
        'yellow': 'objective_1',         # Early Detection
        'darkYellow': 'objective_1',
        'green': 'objective_2',          # Diagnosis & Treatment
        'darkGreen': 'objective_2',
        'cyan': 'objective_3',           # Continuity & Follow-Up
        'blue': 'objective_3',
        'darkBlue': 'objective_3',
        'darkCyan': 'objective_3',
        'magenta': 'unknown_color:magenta',
        'red': 'unknown_color:red',
        'darkRed': 'unknown_color:darkRed',
        'darkMagenta': 'unknown_color:darkMagenta',
    }
    return COLOR_MAP.get(color_lower, f'unknown_color:{color}')


def parse_document(doc_path):
    """Main parsing logic."""
    print(f"\n📄 Parsing: {doc_path}")
    print(f"   File size: {doc_path.stat().st_size / 1024:.0f} KB\n")

    doc = Document(str(doc_path))
    comments = extract_comments(doc_path)

    # Build interview boundaries from Heading 1 elements
    boundaries = build_interview_boundaries(doc)
    print(f"  Found {len(boundaries)} interview boundaries:")
    for _, iid, ptype in boundaries:
        print(f"    {iid} ({ptype})")

    # Parse the raw XML for comment references
    from zipfile import ZipFile
    with ZipFile(doc_path, 'r') as zf:
        document_xml = zf.read('word/document.xml')
    doc_tree = etree.fromstring(document_xml)

    # Build a map of comment ranges: commentRangeStart id → paragraph index
    comment_range_map = {}  # comment_id → paragraph index
    all_paragraphs = doc_tree.findall('.//w:p', NSMAP)

    for idx, para in enumerate(all_paragraphs):
        for crs in para.findall('.//w:commentRangeStart', NSMAP):
            cid = crs.get(f'{{{NSMAP["w"]}}}id')
            if cid:
                comment_range_map[cid] = idx
        for cr in para.findall('.//w:commentReference', NSMAP):
            cid = cr.get(f'{{{NSMAP["w"]}}}id')
            if cid and cid not in comment_range_map:
                comment_range_map[cid] = idx

    print(f"\n  Found {len(comment_range_map)} comment references in body\n")

    # ── Walk paragraphs, extract highlighted segments ──

    findings = []

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # ── Extract highlighted runs from this paragraph ──

        highlighted_segments = []
        current_segment = None

        for run in para.runs:
            run_elem = run._element
            highlight_color = get_highlight_color(run_elem)
            font_color = get_run_color(run_elem)
            run_text = run.text or ''

            effective_color = highlight_color or font_color

            if effective_color:
                if current_segment and current_segment['color'] == effective_color:
                    # Continue the same highlight
                    current_segment['text'] += run_text
                else:
                    # New highlight segment
                    if current_segment:
                        highlighted_segments.append(current_segment)
                    current_segment = {
                        'text': run_text,
                        'color': effective_color,
                        'highlight_color': highlight_color,
                        'font_color': font_color,
                    }
            else:
                if current_segment:
                    # Check if this unhighlighted run is just whitespace between highlighted runs
                    if run_text.strip() == '' and len(run_text) <= 3:
                        current_segment['text'] += run_text
                    else:
                        highlighted_segments.append(current_segment)
                        current_segment = None

        if current_segment:
            highlighted_segments.append(current_segment)

        # ── Process each highlighted segment ──

        for seg in highlighted_segments:
            seg_text = seg['text'].strip()
            if len(seg_text) < 5:  # Skip tiny fragments
                continue

            # Detect facilitator/barrier
            category = detect_facilitator_barrier(seg_text)

            # Clean the text (remove +/- markers)
            clean_text = re.sub(r'\s*\(\+\)\s*$', '', seg_text)
            clean_text = re.sub(r'\s*\(-\)\s*$', '', clean_text)
            clean_text = clean_text.strip()

            # Map color to objective
            objective = map_color_to_objective(seg['color'])

            # Resolve interview from boundaries
            interview_id, participant_type = get_interview_for_paragraph(para_idx, boundaries)

            # Find associated comment
            associated_comment = None
            # Search for comments whose range includes this paragraph
            for cid, pidx in comment_range_map.items():
                if abs(pidx - para_idx) <= 2 and cid in comments:
                    associated_comment = comments[cid]
                    break

            finding = {
                'interview_id': interview_id,
                'participant_type': participant_type,
                'paragraph_index': para_idx,
                'excerpt': clean_text[:300],  # Cap length
                'full_text': clean_text,
                'color': seg['color'],
                'highlight_color': seg['highlight_color'],
                'font_color': seg['font_color'],
                'objective': objective,
                'category': category,  # facilitator | barrier | None
                'comment': associated_comment,
                'theme': associated_comment['text'] if associated_comment else None,
            }
            findings.append(finding)

    return findings


def print_summary(findings):
    """Print a summary of parsed findings."""
    print(f"\n{'='*60}")
    print(f"  PARSING SUMMARY")
    print(f"{'='*60}")
    print(f"  Total findings extracted: {len(findings)}")

    # By interview
    interviews = {}
    for f in findings:
        key = f['interview_id'] or 'Unknown'
        interviews[key] = interviews.get(key, 0) + 1
    print(f"\n  By Interview ({len(interviews)} detected):")
    for iv, count in sorted(interviews.items()):
        print(f"    {iv}: {count} findings")

    # By objective
    objectives = {}
    for f in findings:
        obj = f['objective'] or 'unmapped'
        objectives[obj] = objectives.get(obj, 0) + 1
    print(f"\n  By Objective:")
    for obj, count in sorted(objectives.items()):
        print(f"    {obj}: {count}")

    # By category
    categories = {}
    for f in findings:
        cat = f['category'] or 'unclassified'
        categories[cat] = categories.get(cat, 0) + 1
    print(f"\n  By Category:")
    for cat, count in sorted(categories.items()):
        print(f"    {cat}: {count}")

    # Comments coverage
    with_comments = sum(1 for f in findings if f['comment'])
    print(f"\n  With comments: {with_comments}/{len(findings)} ({100*with_comments/max(len(findings),1):.0f}%)")

    # Colors found
    colors = set(f['color'] for f in findings)
    print(f"\n  Colors found: {sorted(colors)}")

    print(f"{'='*60}\n")


def classify_unclassified(findings):
    """
    Use Anthropic Claude Haiku to classify findings that lack +/− markers
    as facilitator or barrier. Sends in batches of ~40 to stay within limits.
    Returns the number of newly classified findings.
    """
    import os
    import time
    try:
        import anthropic
    except ImportError:
        print("  ⚠ anthropic package not installed — run: pip install anthropic")
        print("    Skipping LLM classification.")
        return 0

    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        # Try reading from .env.local
        env_path = Path(__file__).parent.parent / '.env.local'
        if env_path.exists():
            for line in env_path.read_text().splitlines():
                if line.startswith('ANTHROPIC_API_KEY='):
                    api_key = line.split('=', 1)[1].strip()
                    break
    if not api_key:
        print("  ⚠ ANTHROPIC_API_KEY not found — skipping LLM classification.")
        return 0

    client = anthropic.Anthropic(api_key=api_key)

    SYSTEM_PROMPT = """You are a research analyst classifying findings from breast cancer patient journey interviews.

Each finding is an excerpt from an interview transcript, with an optional human-assigned theme/comment, 
and an objective category:
- objective_1 = Early Detection (screening, symptom recognition, first contact with healthcare)
- objective_2 = Diagnosis & Treatment (diagnostic process, treatment initiation, treatment experience)
- objective_3 = Continuity & Follow-Up (post-treatment care, follow-up adherence, quality of life)

Classify each finding as either:
- "facilitator" — something that HELPED, ENABLED, or POSITIVELY contributed to the patient's journey
- "barrier" — something that HINDERED, DELAYED, or NEGATIVELY impacted the patient's journey

Consider the context carefully. For example:
- "family support during treatment" → facilitator
- "financial burden of chemotherapy" → barrier  
- "early detection through screening" → facilitator
- "delayed diagnosis due to misinterpretation" → barrier
- "psychological counseling available" → facilitator (service availability)
- "fear of diagnosis preventing screening" → barrier

Respond with a JSON array of objects, each with:
- "index": the item index from the input
- "category": "facilitator" or "barrier"
- "confidence": 0.0 to 1.0

Return ONLY valid JSON, no markdown or explanation."""

    # Collect unclassified items
    unclassified = [(i, f) for i, f in enumerate(findings) if f.get('category') is None]
    if not unclassified:
        return 0

    print(f"\n  🤖 Classifying {len(unclassified)} findings via Claude Haiku...")

    BATCH_SIZE = 40
    classified_count = 0
    total_batches = (len(unclassified) + BATCH_SIZE - 1) // BATCH_SIZE

    for batch_num in range(total_batches):
        batch = unclassified[batch_num * BATCH_SIZE : (batch_num + 1) * BATCH_SIZE]
        print(f"    Batch {batch_num + 1}/{total_batches} ({len(batch)} items)...", end=' ', flush=True)

        # Build user message
        items = []
        for idx, (global_idx, f) in enumerate(batch):
            item = {
                'index': global_idx,
                'objective': f.get('objective', 'unknown'),
                'excerpt': f.get('excerpt', '')[:200],
            }
            if f.get('theme'):
                item['theme'] = f['theme']
            items.append(item)

        user_msg = json.dumps(items, ensure_ascii=False)

        try:
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=4096,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_msg}],
            )

            text = response.content[0].text.strip()
            # Strip markdown fences if present
            if text.startswith('```'):
                text = text.split('\n', 1)[1]
                if text.endswith('```'):
                    text = text.rsplit('```', 1)[0]
                text = text.strip()

            results = json.loads(text)

            batch_classified = 0
            for r in results:
                gi = r.get('index')
                cat = r.get('category')
                conf = r.get('confidence', 0)
                if gi is not None and cat in ('facilitator', 'barrier') and gi < len(findings):
                    findings[gi]['category'] = cat
                    findings[gi]['category_source'] = 'llm'
                    findings[gi]['category_confidence'] = conf
                    batch_classified += 1
                    classified_count += 1

            print(f"✓ classified {batch_classified}")

        except Exception as e:
            print(f"✗ error: {e}")

        # Rate limit courtesy
        if batch_num < total_batches - 1:
            time.sleep(0.5)

    return classified_count


def main():
    if not DOC_PATH.exists():
        print(f"❌ Document not found at: {DOC_PATH}")
        print(f"   Place the DOCX file at the project root.")
        sys.exit(1)

    findings = parse_document(DOC_PATH)
    print_summary(findings)

    # Classify unclassified findings using heuristics
    before_unclassified = sum(1 for f in findings if f.get('category') is None)
    newly_classified = classify_unclassified(findings)
    after_unclassified = sum(1 for f in findings if f.get('category') is None)
    print(f"🔍 Heuristic classification:")
    print(f"   Before: {before_unclassified} unclassified")
    print(f"   Classified: {newly_classified}")
    print(f"   Remaining: {after_unclassified} unclassified\n")

    # Show final breakdown
    from collections import Counter
    cats = Counter(f.get('category') or 'unclassified' for f in findings)
    sources = Counter(f.get('category_source', 'marker') for f in findings if f.get('category'))
    print(f"  Final category breakdown:")
    for c, n in sorted(cats.items()):
        print(f"    {c}: {n}")
    print(f"\n  Classification source:")
    for s, n in sorted(sources.items()):
        print(f"    {s}: {n}")
    print()

    # Write JSON output
    output = {
        'parsed_at': datetime.now().isoformat(),
        'source_file': DOC_PATH.name,
        'total_findings': len(findings),
        'findings': findings,
    }

    OUTPUT_PATH.write_text(json.dumps(output, indent=2, ensure_ascii=False), encoding='utf-8')
    print(f"✅ Output written to: {OUTPUT_PATH}")
    print(f"   Review the JSON, then we can merge with LLM-generated clusters.\n")


if __name__ == '__main__':
    main()
